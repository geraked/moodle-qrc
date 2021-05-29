import re
import sys
import subprocess
import os
import platform
import argparse
import pandas as pd
from docx import Document
from docx.shared import Inches, Cm, Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_DIRECTION, WD_TABLE_ALIGNMENT


class QRC:
    def __init__(self, xlsx, docx='', cols=2, rtl=False, open_after=True):
        self.xlsx = xlsx
        self.docx = docx
        self.cols = cols
        self.rtl = rtl
        self.open_after = open_after
        self.run()

    def run(self):
        self.create_document()
        self.write_document()
        self.save_document()
        if self.open_after:
            self.open_document()

    def fix_style(self, name):
        style = self.doc.styles[name]
        style.font.name = 'Calibri'
        if self.rtl:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            style.font.complex_script = True
            style.font.rtl = True

    def create_document(self):
        self.doc = Document()
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(0.5)
            section.right_margin = Cm(0.5)
            section.page_height = Mm(297)
            section.page_width = Mm(210)
            section.header_distance = Cm(0.3)
            section.footer_distance = Cm(0.3)
        self.fix_style('Normal')

    def save_document(self):
        if self.docx == '':
            self.docx = os.path.splitext(self.xlsx)[0] + '.docx'
        self.doc.save(self.docx)

    def open_document(self):
        if platform.system() == 'Darwin':       # macOS
            subprocess.call(('open', self.docx))
        elif platform.system() == 'Windows':    # Windows
            os.startfile(self.docx)
        else:                                   # linux variants
            subprocess.call(('xdg-open', self.docx))

    def clean_text(self, text):
        ctx = str(text)
        ctx = re.sub('[\u00AD\u200F\u00AC]', '\u200C', ctx)
        ctx = re.sub(r'\s?([.:،؛)»؟!?,;])', r'\g<1> ', ctx)
        ctx = re.sub(r'([(«])\s?', r' \g<1>', ctx)
        ctx = re.sub(r'(\s){2,}', r'\g<1>', ctx)  # extra-space
        if self.rtl:
            ctx = ctx.replace('\u064A', '\u06CC')  # ي
            ctx = ctx.replace('\u0643', '\u06A9')  # ك
            ctx = re.sub('\u0020\u0647\u0627' + r'([.:،؛)»؟!?,;]|\s|$)',
                         '\u200C\u0647\u0627' + r'\g<1>', ctx)  # ها
            ctx = re.sub('\u0020\u0647\u0627\u06CC' + r'(\s|$)',
                         '\u200C\u0647\u0627\u06CC' + r'\g<1>', ctx)  # های
            ctx = re.sub('\u0020\u0647\u0627\u06CC\u06CC' + r'(\s|$)',
                         '\u200C\u0647\u0627\u06CC\u06CC' + r'\g<1>', ctx)  # هایی
            ctx = re.sub('\u0647\u0020\u0627\u06CC' + r'([.:،؛)»؟!?,;]|\s|$)',
                         '\u0647\u200C\u0627\u06CC' + r'\g<1>', ctx)  # ای
            ctx = re.sub(r'(\s)' + '\u0645\u06CC\u0020',
                         r'\g<1>' + '\u0645\u06CC\u200C', ctx)  # می
            ctx = re.sub(r'(\s)' + '\u0646\u0645\u06CC\u0020',
                         r'\g<1>' + '\u0646\u0645\u06CC\u200C', ctx)  # نمی
        return ctx.strip()

    def create_table(self):
        c = self.cols * 2
        table = self.doc.add_table(rows=1, cols=c, style='Table Grid')
        if self.rtl:
            table.direction = WD_TABLE_DIRECTION.RTL
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        table.allow_autofit = True
        table._tblPr.xpath(
            "./w:tblW")[0].attrib["{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"] = "auto"
        hdr_cells = table.rows[0].cells
        for cell in hdr_cells:
            cell._tc.tcPr.tcW.type = 'auto'
            cell._tc.tcPr.tcW.w = 0
        ran = (c - 1, -1, -1) if self.rtl else (c,)
        for i in range(*ran):
            con = (i % 2 == 1) if self.rtl else (i % 2 == 0)
            if con:
                hdr_cells[i].text = '#'
            else:
                hdr_cells[i].text = 'ans'
        return table

    def write_document(self):
        df = pd.read_excel(self.xlsx)
        for i in df.index:
            name = df.iloc[i, 0] + ' - ' + df.iloc[i, 1]
            name = self.clean_text(name)
            self.doc.add_heading(name, 2)
            desc = df.iloc[i, 3] + '   -   ' + \
                df.iloc[i, 4] + '   -   ' + df.iloc[i, 5]
            self.doc.add_paragraph(desc)
            table = self.create_table()
            row_cells = []
            k_init = (self.cols * 2 - 1) if self.rtl else 0
            k = k_init
            for j in range(7, len(df.iloc[i])):
                answer = self.clean_text(df.iloc[i, j])
                if k == k_init:
                    row_cells = table.add_row().cells
                    for cell in row_cells:
                        cell._tc.tcPr.tcW.type = 'auto'
                        cell._tc.tcPr.tcW.w = 0
                if self.rtl:
                    row_cells[k].text = str(j - 6)
                    row_cells[k-1].text = answer
                    k -= 2
                    if k < 1:
                        k = k_init
                else:
                    row_cells[k].text = str(j - 6)
                    row_cells[k+1].text = answer
                    k += 2
                    if k > self.cols * 2 - 2:
                        k = k_init
            if (i < len(df.index) - 1):
                self.doc.add_page_break()


def getargs():
    parser = argparse.ArgumentParser(
        description="Convert Moodle .xlsx exam result file to word")
    parser.add_argument('input')
    parser.add_argument('--output', default='')
    parser.add_argument('--cols', type=int, default=2)
    parser.add_argument('--rtl', default=False,
                        action=argparse.BooleanOptionalAction)
    parser.add_argument('--open', default=True,
                        action=argparse.BooleanOptionalAction)
    return(parser.parse_args())


def main():
    args = getargs()
    QRC(xlsx=args.input, docx=args.output,
        cols=args.cols, rtl=args.rtl, open_after=args.open)


if __name__ == '__main__':
    main()
